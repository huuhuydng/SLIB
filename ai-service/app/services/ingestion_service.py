"""
Ingestion Service
Handles document upload, chunking, embedding, and storage in Qdrant
"""

import logging
import io
from typing import List, Dict, Any, Optional
from datetime import datetime

from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.config.settings import get_settings
from app.services.embedding_service import get_embedding_service
from app.services.qdrant_service import get_qdrant_service

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class IngestionService:
    """
    Service for ingesting documents into Qdrant vector database
    
    Workflow:
    1. Load document (PDF, DOCX, or raw text)
    2. Split into chunks using RecursiveCharacterTextSplitter
    3. Generate embeddings using Ollama
    4. Store in Qdrant vector database
    """
    
    def __init__(self):
        settings = get_settings()
        
        # Text splitter configuration
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            length_function=len,
            separators=["\n### ", "\n## ", "\n# ", "\n\n", "\n", ". ", " ", ""]
        )
        
        self.embedding_service = get_embedding_service()
        self.qdrant_service = get_qdrant_service()
        
        logger.info(
            f"[IngestionService] Initialized with chunk_size={settings.chunk_size}, "
            f"overlap={settings.chunk_overlap}, storage=Qdrant"
        )
    
    def ingest_text(
        self, 
        content: str, 
        source: str, 
        category: str = "general",
        metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Ingest raw text content into Qdrant
        
        Args:
            content: Text content to ingest
            source: Source identifier (e.g., filename)
            category: Category for organizing content
            metadata: Additional metadata
            
        Returns:
            Dict with success status and chunk count
        """
        try:
            logger.info(f"[IngestionService] Ingesting text from source: {source}")
            
            # Split text into chunks
            chunks = self.text_splitter.split_text(content)
            logger.info(f"[IngestionService] Created {len(chunks)} chunks")
            
            if not chunks:
                return {
                    "success": False,
                    "message": "No content to ingest",
                    "chunks_created": 0
                }
            
            # Generate embeddings for all chunks
            embeddings = self.embedding_service.embed_texts(chunks)
            logger.info(f"[IngestionService] Generated {len(embeddings)} embeddings")
            
            # Store in Qdrant
            chunks_stored = self.qdrant_service.upsert_vectors(
                vectors=embeddings,
                contents=chunks,
                source=source,
                category=category,
                metadata=metadata
            )
            
            logger.info(f"[IngestionService] Stored {chunks_stored} chunks in Qdrant")
            
            return {
                "success": True,
                "message": f"Successfully ingested {chunks_stored} chunks from {source}",
                "chunks_created": chunks_stored,
                "source": source
            }
                
        except Exception as e:
            logger.error(f"[IngestionService] Error ingesting text: {e}")
            return {
                "success": False,
                "message": f"Error: {str(e)}",
                "chunks_created": 0
            }
    
    def ingest_pdf(
        self, 
        file_content: bytes, 
        filename: str, 
        category: str = "document"
    ) -> Dict[str, Any]:
        """
        Ingest PDF file
        
        Args:
            file_content: PDF file bytes
            filename: Name of the file
            category: Category for the document
            
        Returns:
            Dict with success status and chunk count
        """
        try:
            from PyPDF2 import PdfReader
            
            logger.info(f"[IngestionService] Processing PDF: {filename}")
            
            # Read PDF
            reader = PdfReader(io.BytesIO(file_content))
            text_content = ""
            
            for page in reader.pages:
                text_content += page.extract_text() + "\n\n"
            
            if not text_content.strip():
                return {
                    "success": False,
                    "message": "Could not extract text from PDF",
                    "chunks_created": 0
                }
            
            # Use ingest_text for the rest
            return self.ingest_text(
                content=text_content,
                source=filename,
                category=category,
                metadata={"file_type": "pdf", "pages": len(reader.pages)}
            )
            
        except Exception as e:
            logger.error(f"[IngestionService] Error processing PDF: {e}")
            return {
                "success": False,
                "message": f"PDF processing error: {str(e)}",
                "chunks_created": 0
            }
    
    def ingest_docx(
        self, 
        file_content: bytes, 
        filename: str, 
        category: str = "document"
    ) -> Dict[str, Any]:
        """
        Ingest DOCX file
        
        Args:
            file_content: DOCX file bytes
            filename: Name of the file
            category: Category for the document
            
        Returns:
            Dict with success status and chunk count
        """
        try:
            from docx import Document
            
            logger.info(f"[IngestionService] Processing DOCX: {filename}")
            
            # Read DOCX
            doc = Document(io.BytesIO(file_content))
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                    text_content += "\n"
            
            if not text_content.strip():
                return {
                    "success": False,
                    "message": "Could not extract text from DOCX",
                    "chunks_created": 0
                }
            
            # Use ingest_text for the rest
            return self.ingest_text(
                content=text_content,
                source=filename,
                category=category,
                metadata={"file_type": "docx"}
            )
            
        except Exception as e:
            logger.error(f"[IngestionService] Error processing DOCX: {e}")
            return {
                "success": False,
                "message": f"DOCX processing error: {str(e)}",
                "chunks_created": 0
            }
    
    def get_stats(self) -> Dict[str, Any]:
        """Get statistics about the knowledge base from Qdrant"""
        return self.qdrant_service.get_stats()
    
    def delete_source(self, source: str) -> Dict[str, Any]:
        """Delete all vectors from a specific source"""
        try:
            deleted = self.qdrant_service.delete_by_source(source)
            return {
                "success": True,
                "message": f"Deleted vectors from source: {source}",
                "deleted_count": deleted
            }
        except Exception as e:
            logger.error(f"[IngestionService] Error deleting source: {e}")
            return {
                "success": False,
                "message": f"Error: {str(e)}"
            }


# Singleton instance
_ingestion_service = None


def get_ingestion_service() -> IngestionService:
    """Get singleton IngestionService instance"""
    global _ingestion_service
    if _ingestion_service is None:
        _ingestion_service = IngestionService()
    return _ingestion_service
